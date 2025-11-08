"""
文档加载和处理模块
支持多种文件格式的文档处理和文本分割
"""

import os
import logging
from typing import List, Optional, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
from langchain_core.documents import Document as LangChainDocument
from config.settings import settings

logger = logging.getLogger(__name__)

class DocumentLoader:
    """文档加载器"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
    def load_file(self, file_path: str) -> List[LangChainDocument]:
        """
        加载单个文档文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            分割后的文档块列表
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_extension = file_path.suffix.lower()
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        
        if file_size_mb > settings.MAX_FILE_SIZE_MB:
            raise ValueError(f"文件过大: {file_size_mb:.1f}MB > {settings.MAX_FILE_SIZE_MB}MB")
        
        if file_extension not in settings.SUPPORTED_FILE_TYPES:
            raise ValueError(f"不支持的文件类型: {file_extension}")
        
        try:
            # 提取文本内容
            if file_extension == ".pdf":
                content = self._extract_pdf_text(file_path)
            elif file_extension == ".docx":
                content = self._extract_docx_text(file_path)
            elif file_extension == ".md":
                content = self._extract_markdown_text(file_path)
            elif file_extension == ".txt":
                content = self._extract_text_text(file_path)
            else:
                raise ValueError(f"未实现处理逻辑: {file_extension}")
            
            # 分割文档
            chunks = self.text_splitter.split_text(content)
            
            # 转换为LangChain文档格式
            documents = []
            for i, chunk in enumerate(chunks):
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={
                        "source": str(file_path),
                        "filename": file_path.name,
                        "chunk_id": i,
                        "total_chunks": len(chunks),
                        "file_type": file_extension,
                        "file_size": file_size_mb
                    }
                )
                documents.append(doc)
            
            logger.info(f"处理文件 {file_path.name}: {len(documents)} 个文档块")
            return documents
            
        except Exception as e:
            logger.error(f"处理文件失败 {file_path}: {str(e)}")
            raise
    
    def load_multiple_files(self, file_paths: List[str]) -> List[LangChainDocument]:
        """
        批量处理多个文件
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            所有文档块的列表
        """
        all_documents = []
        
        for file_path in file_paths:
            try:
                documents = self.load_file(file_path)
                all_documents.extend(documents)
            except Exception as e:
                logger.warning(f"跳过文件 {file_path}: {str(e)}")
                continue
        
        logger.info(f"批量处理完成: {len(file_paths)} 个文件, {len(all_documents)} 个文档块")
        return all_documents
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """提取PDF文本"""
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"PDF页面 {page_num} 处理失败: {str(e)}")
                    continue
        
        return "\n".join(text_content)
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """提取DOCX文本"""
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        return "\n".join(text_content)
    
    def _extract_markdown_text(self, file_path: Path) -> str:
        """提取Markdown文本"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_text_text(self, file_path: Path) -> str:
        """提取纯文本"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

class DocumentClassifier:
    """文档分类器 - 支持LLM智能分类"""
    
    def __init__(self, llm_manager=None):
        """
        初始化分类器
        
        Args:
            llm_manager: LLM管理器（可选）。如果提供，则使用LLM智能分类；否则使用规则分类
        """
        self.categories = settings.DOCUMENT_CATEGORIES
        self.priorities = settings.PRIORITY_LEVELS
        self.llm_manager = llm_manager
        self.llm_enabled = llm_manager is not None and hasattr(llm_manager, 'chat_model') and llm_manager.chat_model is not None
        
        if self.llm_enabled:
            logger.info("✅ 文档分类器：使用LLM智能分类模式")
        else:
            logger.info("ℹ️ 文档分类器：使用规则分类模式（简化）")
    
    def classify_document(self, document: LangChainDocument) -> Dict[str, Any]:
        """
        对文档进行分类
        
        Args:
            document: 文档对象
            
        Returns:
            分类结果字典
        """
        content = document.page_content
        
        # 如果启用LLM，优先使用智能分类
        if self.llm_enabled:
            try:
                return self._classify_with_llm(content)
            except Exception as e:
                logger.warning(f"LLM分类失败，回退到规则分类: {e}")
                # 回退到规则分类
        
        # 规则分类（简化模式）
        return self._classify_with_rules(content)
    
    def _classify_with_llm(self, content: str) -> Dict[str, Any]:
        """
        使用LLM进行智能分类
        
        Args:
            content: 文档内容
            
        Returns:
            分类结果字典
        """
        from langchain_core.messages import SystemMessage, HumanMessage
        
        # 限制内容长度（避免token过多）
        analysis_content = content[:1000] if len(content) > 1000 else content
        
        prompt = f"""请分析以下文档内容，并提供分类信息。

文档内容：
{analysis_content}

请按以下格式返回（每行一个字段）：
类别: [从以下选择一个: 工作, 学习, 个人, 参考, 研究, 想法]
优先级: [从以下选择一个: 高, 中, 低]
标签: [提取3-5个关键标签，用逗号分隔]
摘要: [一句话总结，不超过50字]

要求：
1. 类别必须是列表中的一个
2. 标签要具体、简洁（2-4字）
3. 摘要要突出核心主题"""

        messages = [
            SystemMessage(content="你是一个文档分类专家，擅长快速准确地分类和总结文档内容。"),
            HumanMessage(content=prompt)
        ]
        
        try:
            response = self.llm_manager.chat_model.invoke(messages)
            response_text = response.content if hasattr(response, 'content') else str(response)
            
            # 解析LLM响应
            result = self._parse_llm_response(response_text, content)
            logger.info(f"LLM分类完成: 类别={result['category']}, 标签={result['tags']}")
            return result
            
        except Exception as e:
            logger.error(f"LLM调用失败: {e}")
            raise
    
    def _parse_llm_response(self, response_text: str, original_content: str) -> Dict[str, Any]:
        """
        解析LLM返回的分类结果
        
        Args:
            response_text: LLM响应文本
            original_content: 原始文档内容
            
        Returns:
            标准化的分类结果
        """
        lines = response_text.strip().split('\n')
        
        # 默认值
        category = "参考"
        priority = "中"
        tags = []
        summary = original_content[:100] + "..."
        
        # 解析每一行
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if '类别' in line or 'Category' in line.lower():
                # 提取类别
                for cat in self.categories:
                    if cat in line:
                        category = cat
                        break
            
            elif '优先级' in line or 'Priority' in line.lower():
                # 提取优先级
                if '高' in line or 'high' in line.lower():
                    priority = "高"
                elif '低' in line or 'low' in line.lower():
                    priority = "低"
                else:
                    priority = "中"
            
            elif '标签' in line or 'Tag' in line.lower():
                # 提取标签
                tags_part = line.split(':', 1)[-1].strip()
                # 移除可能的方括号
                tags_part = tags_part.strip('[]')
                tags = [t.strip() for t in tags_part.split(',') if t.strip()]
            
            elif '摘要' in line or 'Summary' in line.lower():
                # 提取摘要
                summary_part = line.split(':', 1)[-1].strip()
                if summary_part and len(summary_part) > 10:
                    summary = summary_part
        
        # 验证和修正
        if category not in self.categories:
            category = "参考"
        
        if not tags:
            # 如果LLM没有返回标签，使用规则提取
            tags = self._extract_keywords(original_content)[:5]
        
        return {
            "category": category,
            "priority": priority,
            "summary": summary,
            "tags": ",".join(tags[:5]),  # 最多5个标签
            "confidence": 0.85,  # LLM分类置信度较高
            "classification_method": "llm"  # 标记使用LLM分类
        }
    
    def _classify_with_rules(self, content: str) -> Dict[str, Any]:
        """
        使用规则进行分类（简化模式）
        
        Args:
            content: 文档内容
            
        Returns:
            分类结果字典
        """
        content_lower = content.lower()
        
        category_rules = {
            "工作": ["项目", "会议", "工作", "任务", "计划", "报告", "需求"],
            "学习": ["学习", "教程", "课程", "笔记", "知识", "技能", "练习"],
            "个人": ["日记", "想法", "感悟", "生活", "家庭", "个人", "心情"],
            "参考": ["参考", "文档", "手册", "指南", "规范", "标准", "资料"],
            "研究": ["研究", "分析", "实验", "数据", "结论", "发现", "探索"],
            "想法": ["想法", "创意", "创新", "设计", "概念", "思路", "灵感"]
        }
        
        scores = {}
        for category, keywords in category_rules.items():
            score = sum(1 for keyword in keywords if keyword in content_lower)
            scores[category] = score
        
        # 选择得分最高的类别
        if scores:
            category = max(scores, key=scores.get)
            if scores[category] == 0:
                category = "参考"  # 默认类别
        else:
            category = "参考"
        
        # 根据文件类型和内容长度确定优先级
        if "重要" in content_lower or "urgent" in content_lower:
            priority = "高"
        elif len(content) > 2000:
            priority = "中"
        else:
            priority = "低"
        
        # 生成简短摘要
        summary = content[:100] + "..." if len(content) > 100 else content
        
        # 提取关键词并转换为字符串
        keywords = self._extract_keywords(content)
        
        return {
            "category": category,
            "priority": priority,
            "summary": summary,
            "tags": ",".join(keywords),  # 转换为字符串
            "confidence": round(scores.get(category, 0) / max(len(content.split()) / 100, 1), 3),
            "classification_scores": "; ".join([f"{k}:{v}" for k, v in scores.items() if v > 0]),  # 添加scores作为字符串
            "classification_method": "rules"  # 标记使用规则分类
        }
    
    def _extract_keywords(self, content: str) -> List[str]:
        """提取关键词"""
        # 简单的关键词提取
        # 在实际应用中可以使用更复杂的NLP技术
        common_words = {"的", "了", "是", "在", "有", "和", "与", "或", "但", "而", "这", "那", "个", "等", "及"}
        
        words = content.split()
        word_freq = {}
        
        for word in words:
            word = word.strip('，。！？；：""''()（）[]【】')
            if len(word) > 1 and word not in common_words:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # 返回频率最高的5个词作为标签
        top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:5]
        return [word for word, freq in top_words]

