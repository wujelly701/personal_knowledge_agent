"""
简化版文档加载和处理模块
不使用LangChain的文本分割器，使用基础字符串分割
"""

import os
import logging
import re
from typing import List, Optional, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
from langchain_core.documents import Document as LangChainDocument
from config.settings import settings

logger = logging.getLogger(__name__)

def simple_text_splitter(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """
    简单的文本分割器
    按段落和句子分割，确保语义连贯性
    """
    if not text:
        return []
    
    # 清理文本
    text = text.strip()
    
    # 首先按段落分割
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    
    chunks = []
    current_chunk = ""
    
    for paragraph in paragraphs:
        # 如果当前块加上这个段落超过chunk_size，先保存当前块
        if len(current_chunk) + len(paragraph) + 2 > chunk_size:
            if current_chunk:
                chunks.append(current_chunk)
                # 保持重叠
                overlap_start = max(0, len(current_chunk) - chunk_overlap)
                current_chunk = current_chunk[overlap_start:] + " " + paragraph
            else:
                current_chunk = paragraph
        else:
            if current_chunk:
                current_chunk += "\n\n" + paragraph
            else:
                current_chunk = paragraph
    
    # 添加最后一个块
    if current_chunk:
        chunks.append(current_chunk)
    
    # 如果分段后的块还是太大，进一步分割
    final_chunks = []
    for chunk in chunks:
        if len(chunk) <= chunk_size:
            final_chunks.append(chunk)
        else:
            # 按句子分割
            sentences = re.split(r'[.!?。！？]', chunk)
            current_sentence = ""
            
            for sentence in sentences:
                if len(current_sentence) + len(sentence) + 1 > chunk_size:
                    if current_sentence:
                        final_chunks.append(current_sentence)
                        # 保持重叠
                        overlap_start = max(0, len(current_sentence) - chunk_overlap)
                        current_sentence = current_sentence[overlap_start:] + sentence
                    else:
                        final_chunks.append(sentence)
                else:
                    if current_sentence:
                        current_sentence += "." + sentence
                    else:
                        current_sentence = sentence
            
            if current_sentence:
                final_chunks.append(current_sentence)
    
    return final_chunks

class DocumentLoader:
    """文档加载器"""
    
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        
    def load_file(self, file_path: str) -> List[LangChainDocument]:
        """
        加载单个文档文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            分割后的文档块列表
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_extension = file_path.suffix.lower()
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        
        if file_size_mb > settings.MAX_FILE_SIZE_MB:
            raise ValueError(f"文件过大: {file_size_mb:.1f}MB > {settings.MAX_FILE_SIZE_MB}MB")
        
        if file_extension not in settings.SUPPORTED_FILE_TYPES:
            raise ValueError(f"不支持的文件类型: {file_extension}")
        
        try:
            # 提取文本内容
            if file_extension == ".pdf":
                content = self._extract_pdf_text(file_path)
            elif file_extension == ".docx":
                content = self._extract_docx_text(file_path)
            elif file_extension == ".md":
                content = self._extract_markdown_text(file_path)
            elif file_extension == ".txt":
                content = self._extract_text_text(file_path)
            else:
                supported = ", ".join(settings.SUPPORTED_FILE_TYPES)
                raise ValueError(
                    f"不支持的文件类型: {file_extension}. "
                    f"支持的格式: {supported}"
                )
            
            # 分割文档
            chunks = simple_text_splitter(content, self.chunk_size, self.chunk_overlap)
            
            # 转换为LangChain文档格式
            documents = []
            for i, chunk in enumerate(chunks):
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={
                        "source": str(file_path),
                        "filename": file_path.name,
                        "chunk_id": i,
                        "total_chunks": len(chunks),
                        "file_type": file_extension,
                        "file_size": file_size_mb,
                        "content_hash": str(hash(chunk))  # 添加内容哈希
                    }
                )
                documents.append(doc)
            
            logger.info(f"处理文件 {file_path.name}: {len(documents)} 个文档块")
            return documents
            
        except Exception as e:
            logger.error(f"处理文件失败 {file_path}: {str(e)}")
            raise
    
    def load_multiple_files(self, file_paths: List[str]) -> List[LangChainDocument]:
        """
        批量处理多个文件
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            所有文档块的列表
        """
        all_documents = []
        
        for file_path in file_paths:
            try:
                documents = self.load_file(file_path)
                all_documents.extend(documents)
            except Exception as e:
                logger.warning(f"跳过文件 {file_path}: {str(e)}")
                continue
        
        logger.info(f"批量处理完成: {len(file_paths)} 个文件, {len(all_documents)} 个文档块")
        return all_documents
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """提取PDF文本"""
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"PDF页面 {page_num} 处理失败: {str(e)}")
                    continue
        
        return "\n".join(text_content)
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """提取DOCX文本"""
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        return "\n".join(text_content)
    
    def _extract_markdown_text(self, file_path: Path) -> str:
        """提取Markdown文本"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_text_text(self, file_path: Path) -> str:
        """提取纯文本"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

class DocumentClassifier:
    """文档分类器 - 改进版：多维度分类"""
    
    def __init__(self):
        self.categories = settings.DOCUMENT_CATEGORIES
        self.priorities = settings.PRIORITY_LEVELS
        
        # 增强的分类规则：权重 + 关键词
        self.category_rules = {
            "工作": {
                "keywords": ["项目", "会议", "工作", "任务", "计划", "报告", "需求", "deadline", "团队", "客户"],
                "weight": 1.0,
                "patterns": ["todo", "meeting", "project"]
            },
            "学习": {
                "keywords": ["学习", "教程", "课程", "笔记", "知识", "技能", "练习", "章节", "总结", "复习"],
                "weight": 1.0,
                "patterns": ["python", "java", "编程", "algorithm"]
            },
            "个人": {
                "keywords": ["日记", "想法", "感悟", "生活", "家庭", "个人", "心情", "日程", "schedule"],
                "weight": 0.9,
                "patterns": ["diary", "daily", "schedule"]
            },
            "参考": {
                "keywords": ["参考", "文档", "手册", "指南", "规范", "标准", "资料", "api", "documentation"],
                "weight": 0.8,
                "patterns": ["reference", "guide", "manual"]
            },
            "研究": {
                "keywords": ["研究", "分析", "实验", "数据", "结论", "发现", "探索", "论文", "方法"],
                "weight": 1.0,
                "patterns": ["research", "analysis", "experiment"]
            },
            "想法": {
                "keywords": ["想法", "创意", "创新", "设计", "概念", "思路", "灵感", "brainstorm"],
                "weight": 0.9,
                "patterns": ["idea", "innovation", "creative"]
            }
        }
    
    def classify_document(self, document: LangChainDocument) -> Dict[str, Any]:
        """
        对文档进行多维度分类
        
        Args:
            document: 文档对象
            
        Returns:
            分类结果字典
        """
        content = document.page_content.lower()
        filename = document.metadata.get('filename', '').lower()
        
        # 1. 基于文件名的初步判断
        filename_scores = {}
        for category, rule in self.category_rules.items():
            score = 0
            for pattern in rule['patterns']:
                if pattern in filename:
                    score += 2.0  # 文件名匹配权重更高
            filename_scores[category] = score
        
        # 2. 基于内容的关键词匹配
        content_scores = {}
        for category, rule in self.category_rules.items():
            score = 0
            keywords = rule['keywords']
            weight = rule['weight']
            
            # 计算关键词出现次数
            for keyword in keywords:
                count = content.count(keyword)
                score += count * weight
            
            # 归一化：除以内容长度，避免长文档得分过高
            content_length = len(content.split())
            if content_length > 0:
                score = score / (content_length / 100)  # 每100词归一化
            
            content_scores[category] = score
        
        # 3. 综合得分
        final_scores = {}
        for category in self.category_rules.keys():
            final_scores[category] = filename_scores.get(category, 0) * 0.4 + content_scores.get(category, 0) * 0.6
        
        # 4. 选择得分最高的类别
        if final_scores and max(final_scores.values()) > 0:
            category = max(final_scores, key=final_scores.get)
            confidence = final_scores[category] / (sum(final_scores.values()) + 0.001)  # 避免除0
        else:
            category = "参考"  # 默认类别
            confidence = 0.3
        
        # 5. 根据内容确定优先级
        priority_keywords = {
            "高": ["重要", "urgent", "紧急", "asap", "高优先级", "critical"],
            "中": ["todo", "待办", "计划", "scheduled"],
            "低": []
        }
        
        priority = "低"  # 默认
        for pri, keywords in priority_keywords.items():
            if any(kw in content for kw in keywords):
                priority = pri
                break
        
        # 如果内容很长，默认中优先级
        if priority == "低" and len(content) > 2000:
            priority = "中"
        
        # 6. 生成简短摘要
        summary = content[:100] + "..." if len(content) > 100 else content
        
        # 7. 提取关键词
        keywords = self._extract_keywords(content)
        
        # 8. 将分类得分转换为字符串（ChromaDB metadata 不支持嵌套字典）
        scores_str = "; ".join([f"{k}:{v:.2f}" for k, v in final_scores.items() if v > 0])
        
        return {
            "category": category,
            "priority": priority,
            "summary": summary,
            "tags": ",".join(keywords),
            "confidence": round(confidence, 3),
            "classification_scores": scores_str  # ✅ 字符串格式，如 "工作:5.81; 学习:3.55"
        }
    
    def _extract_keywords(self, content: str) -> List[str]:
        """提取关键词"""
        # 简单的关键词提取
        common_words = {"的", "了", "是", "在", "有", "和", "与", "或", "但", "而"}
        
        words = content.split()
        word_freq = {}
        
        for word in words:
            word = word.strip('，。！？；：""''()（）[]【】')
            if len(word) > 1 and word not in common_words:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # 返回频率最高的5个词作为标签
        top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:5]
        return [word for word, freq in top_words]
